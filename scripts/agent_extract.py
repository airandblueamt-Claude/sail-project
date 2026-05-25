"""
Kind-aware request extractor.

Reads a request document (.docx / .pdf / .txt / email) and writes a draft
gpu_requests row + all child rows the document supports. Two-pass LLM:

  1. classify_kind(text)  -> {kind, confidence}
  2. extract_for_kind(text, kind) -> structured JSON matching that kind

Three kinds map to three different child-table shapes (see Phase 1
schema rebuild — migrations/2026-05-25-gpu-request-rebuild.py):

  new_infra            -> vm_groups + vm_roles + models (optional)
                          + fields (networking + access)
  gpu_allocation       -> models (with min/max count range)
  compute_partnership  -> workloads + phases + contributions
                          + existing_resource_ref

Source is stamped 'agent' and agent_confidence is recorded, along with
the raw LLM JSON for audit. Mahmood / team approves through the existing
UI; this script never edits a decided request.

Usage:
    .venv/bin/python scripts/agent_extract.py <path>                  # dry-run
    .venv/bin/python scripts/agent_extract.py <path> --commit         # write
    .venv/bin/python scripts/agent_extract.py <path> --kind new_infra # override classifier
    .venv/bin/python scripts/agent_extract.py --samples --commit      # process all docs/samples/

Offline / testing without Ollama:
    .venv/bin/python scripts/agent_extract.py --from-json sample.json --commit
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from database import get_db  # noqa: E402

OLLAMA_URL_DEFAULT = "http://10.20.6.61:11434/api/generate"
MODEL_DEFAULT = "nemotron3:33b-q8"
REQUEST_KINDS = ('new_infra', 'gpu_allocation', 'compute_partnership', 'other')


# ── Document readers ───────────────────────────────────────────────────────

def read_document(path: Path) -> str:
    """Flatten any supported doc to plain text. Tables pipe-separated."""
    ext = path.suffix.lower()
    if ext == '.docx':
        return _read_docx(path)
    if ext == '.pdf':
        return _read_pdf(path)
    if ext in ('.txt', '.md', '.eml'):
        return path.read_text(encoding='utf-8', errors='replace')
    raise ValueError(f'Unsupported file type: {ext}')


def _read_docx(path: Path) -> str:
    import docx
    d = docx.Document(path)
    parts: list[str] = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for i, t in enumerate(d.tables):
        parts.append(f"\n[Table {i}]")
        for row in t.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    import pdfplumber
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
            for j, tbl in enumerate(page.extract_tables() or []):
                parts.append(f"\n[Page {i+1} Table {j}]")
                for row in tbl:
                    parts.append(" | ".join((c or "").strip() for c in row))
    return "\n".join(parts)


# ── LLM provider ───────────────────────────────────────────────────────────

@dataclass
class LLMProvider:
    """Minimal interface — one method, returns a parsed JSON dict.

    Today we have one impl (Ollama on 10.20.6.61). Tomorrow this can be
    swapped for Anthropic / OpenAI / a vLLM endpoint without touching the
    extractor logic. The agent_extract caller never sees the transport.
    """
    url: str = OLLAMA_URL_DEFAULT
    model: str = MODEL_DEFAULT
    timeout_s: int = 600
    num_ctx: int = 8192

    def generate(self, prompt: str) -> dict:
        import requests
        r = requests.post(
            self.url,
            json={
                "model": self.model,
                "prompt": prompt,
                "format": "json",
                "stream": False,
                "options": {"temperature": 0.1, "num_ctx": self.num_ctx},
            },
            timeout=self.timeout_s,
        )
        r.raise_for_status()
        raw = r.json().get("response", "")
        try:
            return json.loads(raw)
        except json.JSONDecodeError as e:
            print("LLM returned non-JSON response:", file=sys.stderr)
            print(raw, file=sys.stderr)
            raise SystemExit(f"JSON parse error: {e}")


# ── Classification ─────────────────────────────────────────────────────────

CLASSIFY_PROMPT = """You are an intake assistant for AMT's IT inventory system (SAIL). \
Read the document below and decide what KIND of request it represents. \
Return JSON only.

Kinds:
- "new_infra": Bring-Your-Own-Cloud or full infrastructure brief. \
The document specifies multiple VM groups with vCPU/RAM/disk per role, \
networking requirements, remote-access requirements. May also include \
GPU options. Example: an "OrbitronAI BYOC Deployment Readiness" doc.
- "gpu_allocation": A short list of GPU models the requester needs, \
sometimes with a count range ("2 per module, up to 8"). No VM groups, \
no networking. Example: a forwarded email listing 6 GPU models.
- "compute_partnership": A request for time on existing SAIL hardware, \
structured as workloads (with config + estimated hours) + phases + \
what the requester provides back. Example: ThakaaMed partnership proposal.
- "other": Anything else, or insufficient info to decide.

---DOCUMENT START---
{doc_text}
---DOCUMENT END---

JSON shape:
{{
  "kind": "new_infra" | "gpu_allocation" | "compute_partnership" | "other",
  "confidence": 0.0-1.0,
  "reason": "one short sentence"
}}
"""


def classify_kind(text: str, llm: LLMProvider) -> tuple[str, float, str]:
    out = llm.generate(CLASSIFY_PROMPT.format(doc_text=text[:32000]))
    kind = out.get("kind", "other")
    if kind not in REQUEST_KINDS:
        kind = "other"
    conf = float(out.get("confidence") or 0.0)
    reason = out.get("reason") or ""
    return kind, conf, reason


# ── Extraction prompts (one per kind) ──────────────────────────────────────

EXTRACT_PROMPTS = {
    "new_infra": """Extract the BYOC infrastructure request below into JSON. Return JSON only.

---DOCUMENT START---
{doc_text}
---DOCUMENT END---

JSON shape:
{{
  "title": "short single-line title",
  "use_case": "1-2 sentences on what they want to deploy",
  "requester_org": "organization sending the request, if named",
  "vm_groups": [
    {{
      "name": "Kubernetes Cluster Nodes",
      "summary": "30 VMs recommended",
      "roles": [
        {{"role_name": "Control Plane", "vm_count": 9, "vcpu_per_vm": 4,
          "ram_gb_per_vm": 16, "disk_gb_per_vm": 50,
          "disk_type": "SSD", "os": "Ubuntu 24.04 LTS", "notes": null}}
      ]
    }}
  ],
  "gpu_options": [
    {{"use_case_label": "Up to 14B FP16",
      "model_name": "NVIDIA L4", "vram_gb": 24,
      "gpu_count": 1, "gpu_count_max": null,
      "host_vcpu": 16, "host_ram_gb": 64, "host_disk_gb": 1000,
      "host_os": "Ubuntu 24.04 LTS"}}
  ],
  "networking": {{"subnet": "...", "static_ip": "...", "dns": "...",
                 "tls": "...", "outbound": "...", "notes": "..."}},
  "access":     {{"ssh_vpn": "...", "bastion": "...",
                 "service_account": "...", "notes": "..."}},
  "relationship": {{"wa_ed_investment": "...", "disai_2025": "...",
                   "program": "...", "notes": "..."}},
  "confidence": 0.0-1.0
}}

Rules:
- vm_groups: every group/role with specs in the doc — leave fields null if absent.
- gpu_options: every GPU option listed (they are alternatives, include all).
  In BYOC briefs each GPU row is a full host spec — capture use_case_label
  and host_vcpu/host_ram_gb/host_disk_gb/host_os when the doc gives them.
- networking / access / relationship: only include keys the doc actually
  mentions; empty object if none.
- Be terse. Do not invent fields.
""",

    "gpu_allocation": """Extract the GPU allocation request below into JSON. Return JSON only.

---DOCUMENT START---
{doc_text}
---DOCUMENT END---

JSON shape:
{{
  "title": "short single-line title",
  "use_case": "what they need GPUs for",
  "requester_org": "organization sending the request",
  "gpu_options": [
    {{"model_name": "NVIDIA A100", "vram_gb": 40,
      "gpu_count": 2, "gpu_count_max": 8}}
  ],
  "notes": "anything else worth capturing",
  "confidence": 0.0-1.0
}}

Rules:
- gpu_options: every GPU model listed, in order.
- If the doc says a range like "2 per module, up to 8", set gpu_count=2 and gpu_count_max=8.
- If the doc gives a single count, set gpu_count to it and leave gpu_count_max null.
""",

    "compute_partnership": """Extract the compute-partnership request below into JSON. Return JSON only.

---DOCUMENT START---
{doc_text}
---DOCUMENT END---

JSON shape:
{{
  "title": "short single-line title",
  "use_case": "what they want to use SAIL's existing compute for",
  "requester_org": "organization sending the request",
  "existing_resource_ref": "which SAIL resource the requester wants time on (e.g. 4x A40 cluster)",
  "requested_hours": null,
  "duration_text": "e.g. '12 months (renewable)' or null",
  "workloads": [
    {{"name": "Chest IQ", "config": "2048x2048, batch 8-16, 4-GPU DDP",
      "estimated_hours": 300}}
  ],
  "phases": [
    {{"name": "Phase 1: Infrastructure Setup",
      "target_date": "Jan-26 TBD", "description": "..."}}
  ],
  "contributions": [
    {{"name": "ClearML Integration", "description": "Job scheduling/monitoring",
      "benefit": "Automated resource optimization"}}
  ],
  "relationship": {{"wa_ed_investment": "...", "disai_2025": "...",
                   "program": "...", "notes": "..."}},
  "confidence": 0.0-1.0
}}

Rules:
- workloads: each workload row in the doc — keep config as the doc wrote it.
- phases: each phase / milestone with its target date.
- contributions: things the REQUESTER provides back to SAIL (cluster config,
  knowledge transfer, documentation, etc.).
- relationship: include only keys the doc mentions; empty object if none.
- If the doc gives a target hours like "1,000-2,000 GPU-hours" use the midpoint
  or upper bound as requested_hours and put the full range in notes (later
  reviewers can adjust).
""",

    "other": """Extract whatever request info you can from the document below. Return JSON only.

---DOCUMENT START---
{doc_text}
---DOCUMENT END---

JSON shape:
{{
  "title": "short single-line title",
  "use_case": "1-2 sentences",
  "requester_org": "organization sending the request, if named",
  "notes": "everything else worth capturing in plain text",
  "confidence": 0.0-1.0
}}
""",
}


def extract_for_kind(text: str, kind: str, llm: LLMProvider) -> dict:
    prompt = EXTRACT_PROMPTS[kind].format(doc_text=text[:32000])
    return llm.generate(prompt)


# ── Persistence ────────────────────────────────────────────────────────────

def _next_request_number(conn) -> str:
    year = datetime.now().year
    prefix = f"GPU-{year}-"
    row = conn.execute(
        "SELECT request_number FROM gpu_requests "
        "WHERE request_number LIKE ? ORDER BY id DESC LIMIT 1",
        (f"{prefix}%",),
    ).fetchone()
    if row:
        try:
            n = int(row['request_number'].split('-')[-1]) + 1
        except (ValueError, IndexError):
            n = 1
    else:
        n = 1
    return f"{prefix}{n:04d}"


def _as_int(v: Any) -> int | None:
    if v is None or v == "":
        return None
    try:
        return int(v)
    except (TypeError, ValueError):
        return None


def _as_str(v: Any) -> str | None:
    if v is None:
        return None
    s = str(v).strip()
    return s or None


def persist_request(extracted: dict, kind: str, confidence: float,
                    source_path: Path | None,
                    raw_json: str) -> str:
    """Write the extracted draft into gpu_requests + all child tables relevant to kind."""
    title = _as_str(extracted.get("title")) or (
        source_path.stem if source_path else "Agent-drafted request")
    use_case = _as_str(extracted.get("use_case"))
    requester_org = _as_str(extracted.get("requester_org"))
    notes = _as_str(extracted.get("notes"))
    existing_ref = _as_str(extracted.get("existing_resource_ref"))
    requested_hours = _as_int(extracted.get("requested_hours"))
    duration_text = _as_str(extracted.get("duration_text"))

    with get_db() as conn:
        number = _next_request_number(conn)
        cur = conn.execute(
            """
            INSERT INTO gpu_requests (
                request_number, request_kind, title, use_case,
                requester_name, requester_email, requester_org, requester_type,
                requested_hours, duration_text,
                existing_resource_ref, notes,
                source, agent_confidence, raw_extraction_json
            ) VALUES (?, ?, ?, ?, ?, ?, ?, 'vendor',
                      ?, ?, ?, ?, 'agent', ?, ?)
            """,
            (number, kind, title, use_case,
             requester_org or 'Unknown', '', requester_org,
             requested_hours, duration_text,
             existing_ref, notes, confidence, raw_json),
        )
        req_id = cur.lastrowid

        # gpu_options -> gpu_request_models (used by new_infra + gpu_allocation).
        # In BYOC briefs each row carries full host spec (vCPU/RAM/disk/OS)
        # plus a use_case_label; KFUPM-style short lists leave those null.
        for i, opt in enumerate(extracted.get("gpu_options") or []):
            conn.execute(
                "INSERT INTO gpu_request_models "
                "(request_id, sort_order, use_case_label, model_name, vram_gb, "
                " gpu_count, gpu_count_max, host_vcpu, host_ram_gb, host_disk_gb, host_os) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (req_id, i,
                 _as_str(opt.get("use_case_label")),
                 _as_str(opt.get("model_name")) or "Unknown",
                 _as_int(opt.get("vram_gb")),
                 _as_int(opt.get("gpu_count")),
                 _as_int(opt.get("gpu_count_max")),
                 _as_int(opt.get("host_vcpu")),
                 _as_int(opt.get("host_ram_gb")),
                 _as_int(opt.get("host_disk_gb")),
                 _as_str(opt.get("host_os"))),
            )

        # vm_groups -> gpu_request_vm_groups + gpu_request_vm_roles (new_infra)
        for gi, group in enumerate(extracted.get("vm_groups") or []):
            gname = _as_str(group.get("name"))
            if not gname:
                continue
            cur_g = conn.execute(
                "INSERT INTO gpu_request_vm_groups "
                "(request_id, sort_order, name, summary) VALUES (?, ?, ?, ?)",
                (req_id, gi, gname, _as_str(group.get("summary"))),
            )
            group_id = cur_g.lastrowid
            for ri, role in enumerate(group.get("roles") or []):
                rname = _as_str(role.get("role_name"))
                if not rname:
                    continue
                conn.execute(
                    "INSERT INTO gpu_request_vm_roles "
                    "(group_id, sort_order, role_name, vm_count, vcpu_per_vm, "
                    " ram_gb_per_vm, disk_gb_per_vm, disk_type, os, notes) "
                    "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (group_id, ri, rname,
                     _as_int(role.get("vm_count")),
                     _as_int(role.get("vcpu_per_vm")),
                     _as_int(role.get("ram_gb_per_vm")),
                     _as_int(role.get("disk_gb_per_vm")),
                     _as_str(role.get("disk_type")),
                     _as_str(role.get("os")),
                     _as_str(role.get("notes"))),
                )

        # workloads / phases / contributions (compute_partnership)
        for i, w in enumerate(extracted.get("workloads") or []):
            name = _as_str(w.get("name"))
            if not name:
                continue
            conn.execute(
                "INSERT INTO gpu_request_workloads "
                "(request_id, sort_order, name, config, estimated_hours) "
                "VALUES (?, ?, ?, ?, ?)",
                (req_id, i, name, _as_str(w.get("config")),
                 _as_int(w.get("estimated_hours"))),
            )
        for i, p in enumerate(extracted.get("phases") or []):
            name = _as_str(p.get("name"))
            if not name:
                continue
            conn.execute(
                "INSERT INTO gpu_request_phases "
                "(request_id, sort_order, name, target_date, description) "
                "VALUES (?, ?, ?, ?, ?)",
                (req_id, i, name, _as_str(p.get("target_date")),
                 _as_str(p.get("description"))),
            )
        for i, c in enumerate(extracted.get("contributions") or []):
            name = _as_str(c.get("name"))
            if not name:
                continue
            conn.execute(
                "INSERT INTO gpu_request_contributions "
                "(request_id, sort_order, name, description, benefit) "
                "VALUES (?, ?, ?, ?, ?)",
                (req_id, i, name, _as_str(c.get("description")),
                 _as_str(c.get("benefit"))),
            )

        # networking / access / relationship / document -> gpu_request_fields
        for section in ("networking", "access", "relationship", "document"):
            block = extracted.get(section) or {}
            if isinstance(block, dict):
                for k, v in block.items():
                    val = _as_str(v)
                    if val:
                        conn.execute(
                            "INSERT OR REPLACE INTO gpu_request_fields "
                            "(request_id, section, key, value) VALUES (?, ?, ?, ?)",
                            (req_id, section, str(k), val),
                        )

    return number


# ── CLI ────────────────────────────────────────────────────────────────────

def _process_one(path: Path, llm: LLMProvider, kind_override: str | None,
                 commit: bool) -> None:
    print(f"\n=== {path.name} ===")
    print("[1/3] Reading...")
    text = read_document(path)
    print(f"      {len(text)} chars")

    if kind_override:
        kind, kclass_conf, reason = kind_override, 1.0, "manual override"
        print(f"[2/3] Kind: {kind} (overridden)")
    else:
        print("[2/3] Classifying kind...")
        kind, kclass_conf, reason = classify_kind(text, llm)
        print(f"      kind={kind} confidence={kclass_conf:.2f} reason={reason!r}")

    print(f"[3/3] Extracting with {kind} schema...")
    extracted = extract_for_kind(text, kind, llm)
    confidence = float(extracted.get("confidence") or kclass_conf or 0.0)
    raw_json = json.dumps(extracted, ensure_ascii=False)

    print("      Extracted JSON:")
    for line in json.dumps(extracted, indent=2).splitlines()[:40]:
        print(f"        {line}")
    if len(json.dumps(extracted, indent=2).splitlines()) > 40:
        print("        … (truncated)")

    if not commit:
        print(f"\n      DRY-RUN — pass --commit to write. Would insert kind={kind} "
              f"with confidence={confidence:.2f}.")
        return

    number = persist_request(extracted, kind, confidence, path, raw_json)
    print(f"\n      Inserted as {number} (source=agent, confidence={confidence:.2f})")
    print(f"      View at /gpu/requests/{number}")


def main() -> int:
    ap = argparse.ArgumentParser(description="Agent extractor — turn a request doc into a draft gpu_request.")
    ap.add_argument("doc", type=Path, nargs="?", help="Path to .docx/.pdf/.txt/.eml")
    ap.add_argument("--commit", action="store_true",
                    help="Actually write to sail.db (otherwise dry-run)")
    ap.add_argument("--kind", choices=REQUEST_KINDS,
                    help="Skip the classifier; use this kind directly")
    ap.add_argument("--samples", action="store_true",
                    help=f"Process every file under {REPO_ROOT/'docs/samples'} in one go")
    ap.add_argument("--from-json", type=Path,
                    help="Skip the LLM entirely; load extracted JSON from this file "
                         "and persist it. Requires --kind.")
    ap.add_argument("--url", default=OLLAMA_URL_DEFAULT,
                    help=f"Ollama URL (default: {OLLAMA_URL_DEFAULT})")
    ap.add_argument("--model", default=MODEL_DEFAULT,
                    help=f"Ollama model tag (default: {MODEL_DEFAULT})")
    args = ap.parse_args()

    # --from-json path: skip the LLM entirely. Used for offline testing
    # and to dry-run schema/persistence changes without burning model time.
    if args.from_json:
        if not args.kind:
            sys.exit("--from-json requires --kind to know which child tables to populate.")
        extracted = json.loads(args.from_json.read_text())
        if not args.commit:
            print("DRY-RUN — pass --commit to write the JSON.")
            return 0
        number = persist_request(
            extracted, args.kind,
            float(extracted.get("confidence") or 0.0),
            args.from_json, args.from_json.read_text())
        print(f"Inserted as {number} from {args.from_json.name}")
        return 0

    llm = LLMProvider(url=args.url, model=args.model)

    if args.samples:
        samples_dir = REPO_ROOT / "docs" / "samples"
        for p in sorted(samples_dir.iterdir()):
            if p.suffix.lower() in ('.docx', '.pdf', '.txt', '.md', '.eml'):
                try:
                    _process_one(p, llm, args.kind, args.commit)
                except Exception as e:
                    print(f"      FAILED on {p.name}: {e}")
        return 0

    if not args.doc:
        ap.error("Pass a document path, or --samples, or --from-json")

    if not args.doc.exists():
        sys.exit(f"file not found: {args.doc}")
    _process_one(args.doc, llm, args.kind, args.commit)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
