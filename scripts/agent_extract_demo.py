"""
Demo: take a .docx (the kind Mahmood receives from clients) and turn it into
a draft gpu_requests row via Ollama.

    .venv/bin/python scripts/agent_extract_demo.py docs/samples/OrbitronAI_BYOC_Final_Summary.docx
    .venv/bin/python scripts/agent_extract_demo.py <path> --commit     # actually write to sail.db

Without --commit the script just prints what it would have inserted. This is
deliberately a one-file proof-of-concept: no IMAP poll, no /api/v1 write
endpoint, no UI lane. Goal is to see the extraction shape on a real document
before building the rest of the pipeline.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

import docx
import requests

# Make `from database import get_db` work whether invoked from repo root or scripts/.
REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))
from database import get_db  # noqa: E402

OLLAMA_URL = "http://10.20.6.61:11434/api/generate"
MODEL = "nemotron3:33b-q8"


def extract_text(path: Path) -> str:
    """Flatten the .docx to plain text with table rows pipe-separated."""
    d = docx.Document(path)
    parts: list[str] = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for i, t in enumerate(d.tables):
        parts.append(f"\n[Table {i}]")
        for row in t.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


EXTRACTION_PROMPT = """You are an intake assistant for AMT's IT inventory system (SAIL). \
A client sent the document below to request infrastructure. \
Extract structured fields and return JSON only (no prose, no markdown fences).

---DOCUMENT START---
{doc_text}
---DOCUMENT END---

JSON shape:
{{
  "is_request": true|false,
  "request_kind": "gpu_infra_byoc" | "regular_ticket" | "irrelevant",
  "title": "short single-line title",
  "customer": "organization that is asking",
  "use_case": "1-2 sentences on what they want to deploy",
  "gpu_options": [
    {{"model_name": "...", "vram_gb": 80, "count": 1}}
  ],
  "vm_groups_summary": "plain-text summary of VM groups, one line each",
  "networking_notes": "subnet/IP/DNS/TLS requirements as plain text",
  "access_notes": "ssh/vpn/bastion/service-account requirements as plain text",
  "confidence": 0.0-1.0
}}

Rules:
- gpu_options: include every GPU option the doc lists, even if they are alternatives.
- If the doc is purely informational (no resource request), set is_request=false.
- Be terse. Do not invent fields not in the doc.
"""


def call_ollama(doc_text: str) -> dict:
    prompt = EXTRACTION_PROMPT.format(doc_text=doc_text)
    r = requests.post(
        OLLAMA_URL,
        json={
            "model": MODEL,
            "prompt": prompt,
            "format": "json",
            "stream": False,
            "options": {"temperature": 0.1, "num_ctx": 8192},
        },
        timeout=600,
    )
    r.raise_for_status()
    body = r.json()
    raw = body.get("response", "")
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print("Ollama returned non-JSON response:", file=sys.stderr)
        print(raw, file=sys.stderr)
        raise SystemExit(f"JSON parse error: {e}")


def next_request_number(conn) -> str:
    from datetime import datetime
    year = datetime.now().year
    prefix = f"GPU-{year}-"
    row = conn.execute(
        "SELECT request_number FROM gpu_requests WHERE request_number LIKE ? "
        "ORDER BY id DESC LIMIT 1",
        (f"{prefix}%",),
    ).fetchone()
    n = int(row["request_number"].split("-")[-1]) + 1 if row else 1
    return f"{prefix}{n:04d}"


def insert_request(extracted: dict, source_path: Path) -> str:
    """Write the extracted draft into gpu_requests + gpu_request_models."""
    notes_blocks = []
    if extracted.get("vm_groups_summary"):
        notes_blocks.append("VM groups:\n" + extracted["vm_groups_summary"])
    if extracted.get("networking_notes"):
        notes_blocks.append("Networking:\n" + extracted["networking_notes"])
    if extracted.get("access_notes"):
        notes_blocks.append("Access:\n" + extracted["access_notes"])
    notes_blocks.append(
        f"\n[agent_extract_demo] source={source_path.name} "
        f"confidence={extracted.get('confidence')} "
        f"kind={extracted.get('request_kind')}"
    )
    notes = "\n\n".join(notes_blocks)

    with get_db() as conn:
        number = next_request_number(conn)
        cur = conn.execute(
            """
            INSERT INTO gpu_requests (
                request_number, title, use_case,
                requester_name, requester_email, requester_org, requester_type,
                notes
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                number,
                extracted.get("title") or source_path.stem,
                extracted.get("use_case") or "",
                extracted.get("customer") or "Unknown",
                "",
                extracted.get("customer") or "",
                "vendor",
                notes,
            ),
        )
        req_id = cur.lastrowid
        for i, opt in enumerate(extracted.get("gpu_options") or []):
            conn.execute(
                "INSERT INTO gpu_request_models "
                "(request_id, sort_order, model_name, vram_gb, gpu_count) "
                "VALUES (?, ?, ?, ?, ?)",
                (req_id, i, opt.get("model_name"), opt.get("vram_gb"), opt.get("count")),
            )
    return number


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("doc", type=Path, help="Path to .docx")
    ap.add_argument("--commit", action="store_true",
                    help="Actually write to sail.db (otherwise dry-run)")
    args = ap.parse_args()

    if not args.doc.exists():
        sys.exit(f"file not found: {args.doc}")

    print(f"[1/3] Extracting text from {args.doc.name} ...")
    text = extract_text(args.doc)
    print(f"      {len(text)} chars, {text.count(chr(10))+1} lines")

    print(f"[2/3] Calling Ollama ({MODEL}) ...")
    extracted = call_ollama(text)
    print("      → extracted JSON:")
    print(json.dumps(extracted, indent=2))

    if not args.commit:
        print("\n[3/3] Dry-run — pass --commit to write to sail.db")
        return 0

    if not extracted.get("is_request"):
        print("\n[3/3] is_request=false — skipping insert")
        return 0

    number = insert_request(extracted, args.doc)
    print(f"\n[3/3] Inserted as {number} — view at /gpu/requests/{number}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
